"""Resume file validation and text extraction. Pure functions over in-memory
bytes — no network I/O. `app/main.py` owns everything network-shaped
(Supabase Storage, PostgREST); `app/supabase_client.py` owns that transport.

Story 1.2. Content is inspected by its own bytes, never trusted from a
filename or a client-supplied Content-Type header — a `.pdf` extension
proves nothing. PDFs begin `%PDF-`; DOCX files are ZIP archives beginning
`PK\x03\x04`. See DEV-STATE.md named traps, 2026-07-31.
"""

from __future__ import annotations

import io

import docx
from docx.opc.exceptions import PackageNotFoundError
from pypdf import PdfReader
from pypdf.errors import PdfReadError

# Resumes run roughly 200KB in practice (ARCHITECTURE.md §9's storage
# estimate). 5MB is generous headroom for a text document, not a realistic
# ceiling — it exists only to stop something absurd from being read into
# memory, per the "enforce server-side, before reading the whole file"
# requirement in PHASE-1-SPEC 1.2.
MAX_RESUME_BYTES = 5 * 1024 * 1024

PDF_MAGIC = b"%PDF-"
ZIP_MAGIC = b"PK\x03\x04"

CONTENT_TYPES = {
    "pdf": "application/pdf",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


class ResumeValidationError(ValueError):
    """Raised for any resume that fails type detection or text extraction.
    The message is written for the candidate reading it, not a developer —
    callers surface it as the HTTP 400 body verbatim."""


def detect_file_kind(content: bytes) -> str:
    """Returns "pdf" or "docx" from the file's own bytes, never the filename
    or an extension. Rejects anything else, including a same-shaped-but-wrong
    ZIP (a .xlsx, a plain .zip) — that distinction is made downstream by the
    parser itself, in `_extract_docx_text`, since the ZIP signature alone
    cannot tell a DOCX from any other ZIP container."""
    if content.startswith(PDF_MAGIC):
        return "pdf"
    if content.startswith(ZIP_MAGIC):
        return "docx"
    raise ResumeValidationError("Unsupported file. Please upload a PDF or DOCX resume.")


def extract_text(content: bytes) -> tuple[str, str]:
    """Detects the file kind and extracts its text. Returns (kind, text).

    Never returns an empty or whitespace-only string — raises instead. An
    empty extraction flowing onward would let story 1.3's Resume Analyst
    confidently level a candidate off no evidence at all, so a no-text-layer
    PDF (a scan, an image-only export) must fail loudly here, not silently
    downstream. OCR is out of scope; the candidate is told to re-upload.
    """
    kind = detect_file_kind(content)
    text = _extract_pdf_text(content) if kind == "pdf" else _extract_docx_text(content)

    if not text.strip():
        if kind == "pdf":
            raise ResumeValidationError(
                "This PDF has no extractable text. It looks like a scanned or "
                "image-only document. Please upload a text-based PDF, or export "
                "your resume as a DOCX file instead."
            )
        raise ResumeValidationError(
            "This DOCX file has no extractable text. Please check the file and "
            "upload a resume with real content."
        )
    return kind, text.strip()


def _extract_pdf_text(content: bytes) -> str:
    try:
        reader = PdfReader(io.BytesIO(content))
    except PdfReadError as exc:
        raise ResumeValidationError(
            "Unable to read this PDF. The file may be corrupted. Please "
            "re-export it and try again."
        ) from exc

    if reader.is_encrypted:
        # Some exporters set an empty owner password by default, which pypdf
        # can still open; a real password needs a human, so surface a clear
        # message rather than guessing further.
        try:
            reader.decrypt("")
        except Exception as exc:  # noqa: BLE001 — any decrypt failure means "ask the human"
            raise ResumeValidationError(
                "This PDF is password-protected. Please upload an unprotected "
                "PDF or a DOCX file."
            ) from exc

    try:
        pages_text = [page.extract_text() or "" for page in reader.pages]
    except PdfReadError as exc:
        raise ResumeValidationError(
            "Unable to read this PDF. The file may be corrupted. Please "
            "re-export it and try again."
        ) from exc
    return "\n".join(pages_text)


def _extract_docx_text(content: bytes) -> str:
    try:
        document = docx.Document(io.BytesIO(content))
    except (PackageNotFoundError, KeyError, ValueError) as exc:
        # PackageNotFoundError covers "not a zip at all" or an unreadable
        # package. KeyError/ValueError cover a real zip that IS readable but
        # is not a DOCX package — e.g. missing `[Content_Types].xml` — which
        # is exactly what a same-magic-bytes-but-wrong-format upload (a
        # plain .zip, an .xlsx) produces. Both are "not a valid DOCX" to the
        # candidate; the distinction inside python-docx's own exception
        # hierarchy is not.
        raise ResumeValidationError(
            "This file is not a valid DOCX document. Please upload a PDF or "
            "DOCX resume."
        ) from exc

    parts = [p.text for p in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)
